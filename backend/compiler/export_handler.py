"""
Export handlers for different file formats
Converts TipTap HTML to PDF, Markdown, JSON, CSV, and DOCX
"""
import os
import json
import csv
import tempfile
import logging
from io import StringIO, BytesIO
from html.parser import HTMLParser
from urllib.parse import unquote

try:
    from weasyprint import HTML, CSS
except ImportError:
    HTML = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    Document = None

try:
    from markdownify import markdownify as md
except ImportError:
    md = None

logger = logging.getLogger(__name__)


class TipTapHTMLParser(HTMLParser):
    """Parse TipTap HTML to extract content and metadata"""
    
    def __init__(self):
        super().__init__()
        self.content = []
        self.current_tag = None
        self.math_equations = []
        self.tables = []
        self.current_table = []
        self.current_row = []
        self.in_table = False
        self.in_row = False
        self.in_cell = False
        self.cell_content = []
    
    def handle_starttag(self, tag, attrs):
        attrs_dict = {k: v for k, v in attrs}
        
        if tag == 'span' and 'data-latex' in attrs_dict:
            # Extract LaTeX from data-latex attribute
            latex_encoded = attrs_dict.get('data-latex', '')
            latex = unquote(latex_encoded)
            self.math_equations.append(latex)
            self.content.append(f"$${latex}$$")
        elif tag == 'table':
            self.in_table = True
            self.current_table = []
        elif tag == 'tr' and self.in_table:
            self.in_row = True
            self.current_row = []
        elif tag in ['td', 'th'] and self.in_row:
            self.in_cell = True
            self.cell_content = []
        elif tag == 'h1':
            self.current_tag = 'h1'
        elif tag == 'h2':
            self.current_tag = 'h2'
        elif tag == 'h3':
            self.current_tag = 'h3'
        elif tag == 'p':
            self.current_tag = 'p'
        elif tag == 'li':
            self.current_tag = 'li'
        elif tag == 'pre':
            self.current_tag = 'code'
    
    def handle_endtag(self, tag):
        if tag == 'span' and 'data-latex' in str(self.get_starttag_text() or ''):
            pass  # Already handled in handle_starttag
        elif tag == 'table':
            self.in_table = False
            if self.current_table:
                self.tables.append(self.current_table)
        elif tag == 'tr' and self.in_table:
            self.in_row = False
            if self.current_row:
                self.current_table.append(self.current_row)
        elif tag in ['td', 'th'] and self.in_row:
            self.in_cell = False
            cell_text = ''.join(self.cell_content).strip()
            self.current_row.append(cell_text)
        elif tag in ['h1', 'h2', 'h3', 'p', 'li', 'code']:
            if self.current_tag:
                self.current_tag = None
    
    def handle_data(self, data):
        data = data.strip()
        if data:
            if self.in_cell:
                self.cell_content.append(data)
            elif self.current_tag:
                if self.current_tag == 'li':
                    self.content.append(f"- {data}")
                elif self.current_tag == 'code':
                    self.content.append(f"```\n{data}\n```")
                else:
                    self.content.append(data)
            elif not self.in_table and not self.in_row and not self.in_cell:
                if data:
                    self.content.append(data)


def parse_tiptap_html(html_content):
    """Parse TipTap HTML and extract content"""
    parser = TipTapHTMLParser()
    try:
        parser.feed(html_content)
    except Exception as e:
        logger.warning(f"Error parsing HTML: {e}")
    
    return {
        'content': parser.content,
        'math_equations': parser.math_equations,
        'tables': parser.tables
    }


# ============================================================================
# EXPORT HANDLERS
# ============================================================================

def export_to_pdf(html_content, filename="output.pdf"):
    """
    Export HTML to PDF using WeasyPrint
    
    Args:
        html_content: HTML string to convert
        filename: Output filename
    
    Returns:
        bytes: PDF file content
        
    Raises:
        Exception: If WeasyPrint not installed or conversion fails
    """
    if HTML is None:
        raise ImportError("WeasyPrint not installed. Install with: pip install weasyprint")
    
    try:
        logger.info(f"Exporting to PDF: {filename}")
        
        # Wrap HTML in proper document structure
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1, h2, h3 {{ margin-top: 20px; }}
                table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
                td, th {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                code {{ background-color: #f4f4f4; padding: 2px 5px; }}
                pre {{ background-color: #f4f4f4; padding: 10px; overflow-x: auto; }}
                .math {{ background-color: #f9f9f9; padding: 5px; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        pdf_bytes = HTML(string=full_html).write_pdf()
        logger.info(f"PDF export successful: {len(pdf_bytes)} bytes")
        return pdf_bytes
        
    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        raise


def export_to_markdown(html_content, filename="output.md"):
    """
    Export HTML to Markdown
    
    Args:
        html_content: HTML string to convert
        filename: Output filename
    
    Returns:
        str: Markdown content
    """
    try:
        logger.info(f"Exporting to Markdown: {filename}")
        
        # Use markdownify if available
        if md is not None:
            markdown = md(html_content)
        else:
            # Fallback: simple parsing
            parsed = parse_tiptap_html(html_content)
            markdown = "\n\n".join(parsed['content'])
        
        logger.info(f"Markdown export successful: {len(markdown)} chars")
        return markdown
        
    except Exception as e:
        logger.error(f"Markdown export failed: {e}")
        raise


def export_to_json(html_content, filename="output.json"):
    """
    Export HTML to JSON with structured content
    
    Args:
        html_content: HTML string to convert
        filename: Output filename
    
    Returns:
        str: JSON content
    """
    try:
        logger.info(f"Exporting to JSON: {filename}")
        
        parsed = parse_tiptap_html(html_content)
        
        json_data = {
            "metadata": {
                "format": "tiptap-export",
                "filename": filename,
                "html_size": len(html_content)
            },
            "content": {
                "text": parsed['content'],
                "math_equations": parsed['math_equations'],
                "tables": parsed['tables']
            },
            "statistics": {
                "equation_count": len(parsed['math_equations']),
                "table_count": len(parsed['tables']),
                "content_blocks": len(parsed['content'])
            }
        }
        
        json_str = json.dumps(json_data, indent=2)
        logger.info(f"JSON export successful: {len(json_str)} chars")
        return json_str
        
    except Exception as e:
        logger.error(f"JSON export failed: {e}")
        raise


def export_to_csv(html_content, filename="output.csv"):
    """
    Export HTML tables to CSV
    Extracts all tables from HTML and exports as CSV
    
    Args:
        html_content: HTML string to convert
        filename: Output filename
    
    Returns:
        str: CSV content
    """
    try:
        logger.info(f"Exporting to CSV: {filename}")
        
        parsed = parse_tiptap_html(html_content)
        
        if not parsed['tables']:
            logger.warning("No tables found in HTML")
            return "No tables found in document"
        
        # Convert first table to CSV
        output = StringIO()
        writer = csv.writer(output)
        
        first_table = parsed['tables'][0]
        for row in first_table:
            writer.writerow(row)
        
        csv_str = output.getvalue()
        logger.info(f"CSV export successful: {len(csv_str)} chars")
        return csv_str
        
    except Exception as e:
        logger.error(f"CSV export failed: {e}")
        raise


def export_to_docx(html_content, filename="output.docx"):
    """
    Export HTML to DOCX (Microsoft Word)
    
    Args:
        html_content: HTML string to convert
        filename: Output filename
    
    Returns:
        bytes: DOCX file content
        
    Raises:
        Exception: If python-docx not installed or conversion fails
    """
    if Document is None:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    try:
        logger.info(f"Exporting to DOCX: {filename}")
        
        parsed = parse_tiptap_html(html_content)
        doc = Document()
        
        # Add content
        for item in parsed['content']:
            if item.startswith('# '):
                doc.add_heading(item[2:], level=1)
            elif item.startswith('## '):
                doc.add_heading(item[3:], level=2)
            elif item.startswith('### '):
                doc.add_heading(item[4:], level=3)
            elif item.startswith('- '):
                doc.add_paragraph(item[2:], style='List Bullet')
            elif item.startswith('$$'):
                # Add math as paragraph (Word doesn't natively support LaTeX)
                p = doc.add_paragraph()
                p.add_run(f"[Math: {item[2:-2]}]").italic = True
            elif item.startswith('```'):
                doc.add_paragraph(item[3:-3], style='Normal').font.name = 'Courier New'
            else:
                doc.add_paragraph(item)
        
        # Add tables
        for table_data in parsed['tables']:
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = str(cell_data)
        
        # Save to bytes
        bytes_io = BytesIO()
        doc.save(bytes_io)
        docx_bytes = bytes_io.getvalue()
        
        logger.info(f"DOCX export successful: {len(docx_bytes)} bytes")
        return docx_bytes
        
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        raise


# ============================================================================
# EXPORT DISPATCHER
# ============================================================================

EXPORT_HANDLERS = {
    'pdf': export_to_pdf,
    'md': export_to_markdown,
    'markdown': export_to_markdown,
    'json': export_to_json,
    'csv': export_to_csv,
    'docx': export_to_docx,
}


def export_html(html_content, export_format, filename="output"):
    """
    Main export dispatcher - routes to appropriate handler
    
    Args:
        html_content: HTML to export
        export_format: Format (pdf, md, json, csv, docx)
        filename: Base filename without extension
    
    Returns:
        tuple: (content, content_type, file_extension)
               - content: bytes or str
               - content_type: MIME type
               - file_extension: file extension
    
    Raises:
        ValueError: If unsupported format
        Exception: If export fails
    """
    format_lower = export_format.lower().strip()
    
    if format_lower not in EXPORT_HANDLERS:
        raise ValueError(f"Unsupported export format: {format_lower}. Supported: {list(EXPORT_HANDLERS.keys())}")
    
    handler = EXPORT_HANDLERS[format_lower]
    
    try:
        if format_lower == 'pdf':
            content = export_to_pdf(html_content, filename)
            return (content, 'application/pdf', 'pdf')
        
        elif format_lower in ['md', 'markdown']:
            content = export_to_markdown(html_content, filename)
            return (content, 'text/markdown', 'md')
        
        elif format_lower == 'json':
            content = export_to_json(html_content, filename)
            return (content, 'application/json', 'json')
        
        elif format_lower == 'csv':
            content = export_to_csv(html_content, filename)
            return (content, 'text/csv', 'csv')
        
        elif format_lower == 'docx':
            content = export_to_docx(html_content, filename)
            return (content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'docx')
    
    except Exception as e:
        logger.error(f"Export error for format {format_lower}: {e}")
        raise
